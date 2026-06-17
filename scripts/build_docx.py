"""Génère l'article au format Word (preprint/article_fr.docx).

Le texte est révisé pour un style sobre et fluide. Le sourçage est explicite
(citations entre parenthèses, références en fin de document, dont le lien vers
l'historique des auteurs pour le contenu repris de Wikipédia, CC BY-SA).

Usage :
    python scripts/build_docx.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "preprint" / "article_fr.docx"

TITLE = "À l'échelle d'un continent, l'open data n'est pas uniforme"
SUBTITLE = (
    "Reconstituer une donnée manquante avec un pipeline imparfait mais utile : "
    "l'exemple des équipements sportifs en Afrique via OpenStreetMap (2026)"
)
BYLINE = "Benoît Prieur (thepriben)"
PREPRINT_LINE = "Preprint — instantané OpenStreetMap du 15 juin 2026"

# (label, count) pour le tableau de distribution.
DISTRIBUTION = [
    ("soccer", "Football", "61 679"),
    ("tennis", "Tennis", "26 107"),
    ("basketball", "Basketball", "10 774"),
    ("multi", "Multi-sport / polyvalent", "7 557"),
    ("netball", "Netball", "3 821"),
    ("golf", "Golf", "1 835"),
    ("volleyball", "Volley-ball", "1 786"),
    ("cricket", "Cricket", "1 484"),
    ("handball", "Handball", "1 290"),
    ("bowls", "Boulingrin (bowls)", "1 238"),
    ("field_hockey", "Hockey sur gazon", "849"),
    ("rugby_union", "Rugby à XV", "827"),
    ("equestrian", "Équitation", "809"),
    ("athletics", "Athlétisme", "757"),
    ("padel", "Padel", "552"),
]

REFERENCES = [
    "Bale, J. & Cronin, M. (dir.) (2003). Sport and Postcolonialism. Berg.",
    "Geofabrik. Africa : OpenStreetMap Data Extracts. "
    "https://download.geofabrik.de/africa.html",
    "Goodchild, M. F. (2007). « Citizens as sensors: the world of volunteered "
    "geography ». GeoJournal, 69(4), 211-221.",
    "Haklay, M. & Weber, P. (2008). « OpenStreetMap: User-Generated Street Maps ». "
    "IEEE Pervasive Computing, 7(4), 12-18.",
    "Neis, P. & Zielstra, D. (2014). « Recent Developments and Future Trends in "
    "Volunteered Geographic Information Research: The Case of OpenStreetMap ». "
    "Future Internet, 6(1), 76-106.",
    "OpenStreetMap Foundation. About OpenStreetMap. "
    "https://www.openstreetmap.org/about",
    "OpenStreetMap Wiki. Key:sport. https://wiki.openstreetmap.org/wiki/Key:sport",
    "OpenStreetMap. Open Database License (ODbL). "
    "https://www.openstreetmap.org/copyright",
    "Osmium Tool. https://osmcode.org/osmium-tool/",
    "Wikipédia. « Netball », auteurs et historique (consulté en juin 2026), "
    "sous licence CC BY-SA. "
    "https://fr.wikipedia.org/w/index.php?title=Netball&action=history",
    "World Netball. History of the game. https://netball.sport/",
]


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15
    return p


def add_lead(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(label + " ")
    run.bold = True
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)


def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # --- Titre ---
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(20)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = st.add_run(SUBTITLE)
    rs.italic = True
    rs.font.size = Pt(12)
    rs.font.color.rgb = RGBColor(0x55, 0x5F, 0x6E)

    by = doc.add_paragraph()
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by.add_run(BYLINE).bold = True
    pr = doc.add_paragraph()
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr.add_run(PREPRINT_LINE).font.size = Pt(10)

    # --- Résumé ---
    doc.add_heading("Résumé", level=1)
    add_lead(
        doc, "Contexte.",
        "De nombreux objets d'intérêt public ne disposent d'aucune donnée "
        "ouverte fiable, consolidée et comparable. Dans ces situations, une "
        "posture possible consiste non pas à attendre une base de référence "
        "inexistante, mais à produire soi-même une donnée exploitable à partir "
        "de sources citoyennes, en explicitant ses limites et en valorisant ses "
        "avantages. Les équipements sportifs en Afrique constituent un cas "
        "d'école : aucun recensement continental ouvert n'existe, alors même que "
        "leur répartition éclaire des dynamiques territoriales, des héritages "
        "historiques et des inégalités d'accès.",
    )
    add_lead(
        doc, "Objectifs.",
        "Cet article applique cette démarche générale au cas africain. Il teste "
        "l'hypothèse qu'une base citoyenne mondiale, OpenStreetMap (OSM), peut "
        "servir de proxy exploitable. L'enjeu est double : produire une première "
        "mesure du volume, de la répartition et de la typologie des équipements "
        "sportifs, et montrer qu'une telle extraction, parce qu'elle est "
        "reproductible, peut être rejouée périodiquement pour suivre des "
        "évolutions.",
    )
    add_lead(
        doc, "Méthodes.",
        "À partir de l'extrait Afrique d'OSM (Geofabrik, version du 15 juin "
        "2026), une chaîne automatisée filtre, géolocalise et nettoie plus de "
        "123 000 sites sportifs à l'aide d'osmium, de Python et de Folium. Le "
        "pipeline ne prétend pas à l'exhaustivité : il vise à apporter des "
        "réponses utiles là où il n'en existait aucune.",
    )
    add_lead(
        doc, "Résultats.",
        "Le football domine très largement, mais des disciplines discrètes comme "
        "le netball ou le boulingrin (bowls) apparaissent étonnamment présentes "
        "dans certains territoires. La répartition de plusieurs sports (cricket, "
        "rugby, hockey sur gazon) recoupe nettement les anciennes aires "
        "coloniales britanniques, ces disciplines étant quasi absentes des "
        "ex-colonies francophones et lusophones.",
    )
    add_lead(
        doc, "Conclusion.",
        "Imparfaite et non représentative, la donnée OSM constitue néanmoins un "
        "point d'entrée robuste. Sa reproductibilité ouvre la voie à un "
        "observatoire annualisable du fait sportif africain, en passant d'un "
        "instantané isolé à un suivi dans le temps.",
    )
    kw = doc.add_paragraph()
    kw.add_run("Mots-clés : ").bold = True
    kw.add_run(
        "OpenStreetMap, open data, sport, Afrique, observatoire reproductible, "
        "netball, Commonwealth."
    )

    # --- 1. Introduction ---
    doc.add_heading("1. Introduction : produire la donnée plutôt que l'attendre", level=1)
    add_body(
        doc,
        "Le point de départ de ce travail n'est pas une base de données, mais "
        "une situation très courante : l'absence de donnée. Dans bien des "
        "domaines, il n'existe aucune source ouverte, fiable, harmonisée et "
        "comparable, et l'enquête s'arrête souvent là. On peut pourtant adopter "
        "une autre posture, plus constructive : fabriquer une donnée à partir de "
        "matériaux déjà disponibles, en particulier les sources citoyennes, à "
        "condition d'en expliciter rigoureusement les limites. Une telle donnée "
        "n'est jamais « la » référence manquante ; elle en est une approximation "
        "utile, dotée de limites assumées (couverture partielle, biais de "
        "contribution, métadonnées hétérogènes) mais aussi d'avantages propres : "
        "elle est ouverte, géolocalisée, mondiale, et surtout reproductible.",
    )
    add_body(
        doc,
        "Le présent article décline cette démarche générale sur un cas "
        "particulièrement net : les équipements sportifs en Afrique. Il n'existe "
        "pas, à l'échelle du continent, de recensement officiel, harmonisé et "
        "ouvert de ces infrastructures. Les statistiques, lorsqu'elles existent, "
        "sont nationales, éparses, rarement géolocalisées, souvent inaccessibles. "
        "Cette lacune est en soi révélatrice de la faible institutionnalisation "
        "de la donnée sportive et, plus largement, des angles morts de l'open "
        "data.",
    )
    add_body(
        doc,
        "Face à ce vide, la stratégie retenue est pragmatique : utiliser ce qui "
        "existe déjà. OpenStreetMap (OSM), projet de cartographie collaborative "
        "lancé en 2004 (OpenStreetMap Foundation), agrège les contributions de "
        "centaines de milliers de bénévoles et constitue l'exemple le plus abouti "
        "d'information géographique volontaire (Goodchild, 2007 ; Haklay & Weber, "
        "2008). Les contributeurs y décrivent terrains, stades et pitches à "
        "l'aide de tags normalisés, ce qui autorise des extractions thématiques à "
        "grande échelle.",
    )
    add_body(
        doc,
        "Les limites de l'approche sont assumées d'emblée. La donnée OSM n'est ni "
        "exhaustive ni représentative : un terrain de football informel, très "
        "répandu en milieu urbain comme rural, n'est pas toujours cartographié "
        "comme tel, et la densité de contribution varie fortement d'un pays à "
        "l'autre (Neis & Zielstra, 2014). En contrepartie, les avantages sont "
        "réels, et c'est là tout l'intérêt de la démonstration :",
    )
    add_bullets(
        doc,
        [
            "obtenir une première mesure (un ordre de grandeur géolocalisé) là où "
            "il n'existait aucun chiffre ;",
            "accéder à des lectures inédites, impossibles sans donnée "
            "spatialisée : la distribution des disciplines révèle par exemple la "
            "persistance des sports « britanniques » (cricket, rugby, hockey sur "
            "gazon) dans les territoires des anciennes colonies du Commonwealth, "
            "motif qu'aucune statistique agrégée ne ferait apparaître ;",
            "poser les bases d'un observatoire reproductible : la chaîne étant "
            "entièrement automatisée, elle peut être relancée chaque année, ou "
            "appliquée à des dumps anciens (2015, 2020, …), pour mesurer des "
            "évolutions plutôt qu'un seul instant.",
        ],
    )
    add_body(
        doc,
        "L'analyse porte ainsi sur trois dimensions : le volume et la répartition "
        "géographique des équipements recensés en 2026 ; leur distribution par "
        "discipline ; et l'exploration d'un héritage, la persistance des sports "
        "« britanniques » dans les territoires historiquement liés au "
        "Commonwealth, exemple concret de la valeur ajoutée d'une donnée que l'on "
        "a soi-même construite.",
    )

    # --- 2. Méthode ---
    doc.add_heading("2. Approche méthodologique", level=1)
    doc.add_heading("2.1. OpenStreetMap comme source d'open data", level=2)
    add_body(
        doc,
        "OSM est plus qu'une alternative ouverte aux systèmes cartographiques "
        "institutionnels : c'est une œuvre collective fondée sur la logique du "
        "terrain. Les objets (routes, bâtiments, lieux d'activité, "
        "infrastructures sportives) y sont géolocalisés et décrits par des "
        "balises normalisées. Pour les équipements sportifs, les contributeurs "
        "mobilisent principalement leisure=pitch, leisure=stadium et "
        "building=stadium, le tag sport=* précisant la ou les disciplines "
        "pratiquées (OpenStreetMap Wiki, Key:sport).",
    )
    add_body(
        doc,
        "Cette dimension participative est précisément ce qui rend OSM pertinent "
        "là où la donnée officielle fait défaut : des espaces sportifs ignorés "
        "des bases institutionnelles peuvent être repérés et partagés par des "
        "acteurs locaux. La contrepartie est connue : couverture inégale, qualité "
        "variable des métadonnées, dénominations locales floues, approximations "
        "de bonne foi. D'où la nécessité d'un traitement critique et d'une étape "
        "de nettoyage explicite. L'ensemble des données dérivées reste soumis à "
        "l'Open Database License (OpenStreetMap, ODbL).",
    )
    doc.add_heading("2.2. Construction et nettoyage du jeu de données", level=2)
    add_body(
        doc,
        "L'extraction part de l'extrait continental au format .osm.pbf mis à "
        "disposition par Geofabrik, dans sa version du 15 juin 2026 (Geofabrik, "
        "Africa). Le traitement se déroule en plusieurs étapes (détaillées en "
        "annexe) :",
    )
    add_bullets(
        doc,
        [
            "Filtrage des objets portant leisure=pitch, leisure=stadium ou "
            "building=stadium à l'aide de l'outil en ligne de commande osmium "
            "(Osmium Tool).",
            "Conversion au format GeoJSON, puis extraction de la géométrie et du "
            "tag sport avec les bibliothèques Python pandas et shapely (les "
            "géométries linéaires et surfaciques sont ramenées à un point "
            "représentatif : milieu de ligne ou centroïde).",
            "Décomposition des déclarations multiples : un même site peut porter "
            "plusieurs disciplines (p. ex. soccer;basketball). Le séparateur "
            "officiel « ; » (OpenStreetMap Wiki, Key:sport) sert à produire un "
            "enregistrement par sport.",
            "Normalisation de la longue traîne : harmonisation de la casse et des "
            "espaces, correction des fautes fréquentes (socer → soccer, "
            "volley-ball → volleyball, athletisme → athletics) et regroupement des "
            "variantes. Le tag générique « football », déconseillé par "
            "OpenStreetMap au profit de « soccer », est rattaché à ce dernier. "
            "Cette étape ramène le nombre de tags distincts d'environ 250 à 194.",
        ],
    )
    add_body(
        doc,
        "Le corpus final comporte 123 936 sites et 126 928 occurrences sport (un "
        "site pouvant en cumuler plusieurs). Le brut est conservé intact "
        "(data/sports_facilities.csv) afin que tout choix de nettoyage reste "
        "auditable et rejouable.",
    )
    doc.add_heading("2.3. Une chaîne pensée pour l'évolution", level=2)
    add_body(
        doc,
        "Le principal atout de cette méthode n'est pas le chiffre de 2026, mais "
        "sa reproductibilité. La même chaîne, appliquée à l'extrait de l'année "
        "suivante, produit une mesure directement comparable ; appliquée aux "
        "archives d'OSM (history dumps), elle reconstitue rétrospectivement une "
        "trajectoire. On passe alors d'un instantané isolé à une série "
        "temporelle : on peut suivre la densification de la contribution, "
        "l'apparition de nouvelles disciplines (padel, pickleball) ou la "
        "dynamique d'équipement de certaines régions.",
    )

    # --- 3. Volume et distribution ---
    doc.add_heading("3. Volume et distribution des équipements", level=1)
    doc.add_heading("3.1. Une géographie de la concentration", level=2)
    add_body(
        doc,
        "Chaque site étant géolocalisé, une carte de densité a été produite avec "
        "Folium. Une première lecture fait apparaître :",
    )
    add_bullets(
        doc,
        [
            "une concentration nette autour des capitales et métropoles "
            "régionales (Lagos, Johannesburg, Nairobi, Le Caire, Accra…) ;",
            "des disparités marquées entre littoraux et zones sahéliennes ou "
            "enclavées ;",
            "des points isolés dans des territoires faiblement couverts, qui "
            "traduisent autant la réalité de l'équipement que l'inégale activité "
            "des communautés OSM locales.",
        ],
    )
    doc.add_heading(
        "3.2. Distribution par discipline et prédominance du football", level=2
    )
    add_body(
        doc,
        "L'explosion des sports multiples puis le comptage par discipline donnent "
        "la hiérarchie suivante (principaux tags, instantané 2026) :",
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ("Tag OSM", "Sport (interprété)", "Nombre de sites")):
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
    for tag, label, count in DISTRIBUTION:
        cells = table.add_row().cells
        cells[0].text = tag
        cells[1].text = label
        cells[2].text = count
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    add_body(
        doc,
        "Le football concentre à lui seul près de la moitié des équipements "
        "étiquetés. Suivent le tennis, le basketball, puis, après la catégorie "
        "générique « multi-sport », le netball, dont le rang (5e) constitue le "
        "premier résultat réellement surprenant et appelle une lecture "
        "culturelle.",
    )

    # --- 4. Lecture culturelle ---
    doc.add_heading("4. Lecture culturelle et géopolitique", level=1)
    add_body(
        doc,
        "Au-delà du comptage, la dimension spatiale autorise des hypothèses "
        "géopolitiques. Certaines disciplines, marginales dans les grands récits "
        "sportifs internationaux, se révèlent fortement implantées dans des "
        "espaces précis, souvent corrélés à une histoire coloniale partagée.",
    )
    doc.add_heading("4.1. Le netball : un sport mineur, un rang majeur", level=2)
    add_body(
        doc,
        "Le netball se classe cinquième (3 821 sites), devant des sports mondiaux "
        "comme le volley-ball, le cricket ou le rugby. Issu d'une adaptation du "
        "basketball dans le contexte éducatif britannique du début du XXe siècle, "
        "il est aujourd'hui particulièrement pratiqué dans les pays anglophones "
        "du Commonwealth, souvent par les femmes et en milieu scolaire (World "
        "Netball ; Wikipédia, « Netball »).",
    )
    add_body(
        doc,
        "La carte dédiée révèle une concentration en Afrique de l'Est et australe "
        "(Kenya, Afrique du Sud, Malawi, Zimbabwe), précisément les espaces où "
        "l'héritage colonial britannique a intégré ce sport aux politiques "
        "d'éducation physique. À l'inverse, les pays d'héritage francophone ou "
        "lusophone n'en présentent quasiment aucune occurrence. Le netball "
        "illustre ainsi comment une discipline peu médiatisée peut occuper un "
        "rôle territorial structurant et fonctionner comme marqueur identitaire.",
    )
    doc.add_heading(
        "4.2. Sports « britanniques » et territoires du Commonwealth", level=2
    )
    add_body(
        doc,
        "D'autres disciplines historiquement liées à la culture britannique, à "
        "savoir le cricket, le rugby à XV (rugby_union), le hockey sur gazon "
        "(field_hockey) et le boulingrin (bowls), confirment cette géographie du "
        "sport. Croiser leurs localisations avec les anciennes aires coloniales "
        "fait apparaître des motifs nets :",
    )
    add_bullets(
        doc,
        [
            "le rugby, pourtant populaire en France, est très peu présent dans "
            "les pays francophones sur OSM, mais bien implanté en Afrique du Sud, "
            "Namibie, Kenya, Zimbabwe ou Ouganda, qui concentrent l'essentiel des "
            "827 occurrences de rugby_union ;",
            "le cricket suit une logique encore plus marquée, concentré en "
            "Afrique du Sud, au Kenya, au Nigéria et au Ghana, et quasi inexistant "
            "dans les États francophones ;",
            "le hockey sur gazon, emblématique des pays du Commonwealth, se "
            "déploie selon une distribution comparable ;",
            "le boulingrin (bowls) reste très spécifique de l'Afrique australe "
            "anglophone.",
        ],
    )
    add_body(
        doc,
        "La carte thématique associe une couleur à chaque discipline (hockey sur "
        "gazon en rouge, cricket en vert, rugby à XV en bleu, boulingrin en "
        "violet) et matérialise les zones d'héritage colonial britannique. Loin "
        "d'être de simples lieux d'exercice, les équipements traduisent des "
        "histoires longues, des systèmes éducatifs différenciés et des formes "
        "durables d'hégémonie culturelle. Le sport apparaît comme un vecteur "
        "silencieux mais structurant des territorialités post-impériales (sur le "
        "sport comme instrument colonial, voir Bale & Cronin, 2003).",
    )

    # --- 5. Limites ---
    doc.add_heading("5. Limites", level=1)
    add_body(
        doc,
        "La portée de ces résultats est bornée par la nature même de la source. "
        "La couverture OSM est inégale et corrélée au développement des "
        "communautés locales (Neis & Zielstra, 2014) : un faible nombre "
        "d'équipements peut signaler un défaut de cartographie autant qu'un "
        "défaut d'infrastructure. Les terrains informels échappent largement au "
        "recensement. Enfin, le tag sport est renseigné de façon hétérogène, et "
        "la catégorie multi masque une part de l'information. Ces biais invitent "
        "à lire les cartes comme des indices de structure plutôt que comme des "
        "dénombrements absolus, ce qui n'enlève rien à leur intérêt comparatif et "
        "évolutif.",
    )

    # --- Conclusion ---
    doc.add_heading("Conclusion", level=1)
    add_body(
        doc,
        "Ce travail propose une première lecture géographique des équipements "
        "sportifs en Afrique là où aucune donnée ouverte ne l'autorisait. Trois "
        "constats se dégagent : le football domine très largement l'espace "
        "sportif cartographié ; des disciplines peu visibles, comme le netball, "
        "témoignent d'ancrages territoriaux hérités ; et le croisement entre "
        "aires culturelles et sports d'origine britannique dessine une géographie "
        "post-impériale du sport.",
    )
    add_body(
        doc,
        "Surtout, l'étude vaut comme démonstration de méthode. Un pipeline "
        "imparfait, mais entièrement reproductible, suffit à apporter des "
        "réponses concrètes : il ne fige pas un état, il ouvre la possibilité "
        "d'un observatoire annualisable, comparable dans le temps et croisable "
        "avec d'autres variables (démographiques, linguistiques, climatiques, "
        "politiques). C'est dans cette capacité à mesurer le mouvement, plus que "
        "dans le chiffre de 2026, que réside l'intérêt de mobiliser l'open data "
        "citoyen pour combler un vide informationnel.",
    )

    # --- Annexe ---
    doc.add_heading("Annexe méthodologique", level=1)
    add_body(doc, "1. Téléchargement de l'extrait Afrique (Geofabrik)")
    add_code(
        doc,
        [
            "wget -c -O africa-latest.osm.pbf \\",
            "  https://download.geofabrik.de/africa-latest.osm.pbf",
        ],
    )
    add_body(doc, "2. Filtrage des entités pertinentes (osmium)")
    add_code(
        doc,
        [
            "osmium tags-filter africa-latest.osm.pbf \\",
            "  nwr/leisure=pitch nwr/leisure=stadium nwr/building=stadium \\",
            "  --overwrite -o filtered_data.osm",
        ],
    )
    add_body(doc, "3. Conversion en GeoJSON")
    add_code(
        doc,
        ["osmium export filtered_data.osm -f geojson --overwrite -o temp.geojson"],
    )
    add_body(doc, "4. Nettoyage, comptage et cartographie (voir scripts/ du dépôt)")
    add_code(
        doc,
        [
            "python scripts/build_counts.py",
            "python scripts/make_map.py heatmap",
            "python scripts/make_map.py sport --sport netball "
            '--color "#1d3557" --name netball',
            "python scripts/make_map.py british",
        ],
    )

    # --- Références ---
    doc.add_heading("Références", level=1)
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)

    # --- Crédits ---
    doc.add_heading("Crédits", level=1)
    add_body(
        doc,
        "Données : © les contributeurs d'OpenStreetMap, sous licence ODbL. "
        "Extrait continental fourni par Geofabrik. Le passage sur le netball "
        "reprend des éléments de l'article « Netball » de Wikipédia (auteurs et "
        "historique : https://fr.wikipedia.org/w/index.php?title=Netball&action=history), "
        "sous licence CC BY-SA. Code et figures : Benoît Prieur (thepriben).",
    )

    # --- Propriétés ---
    doc.core_properties.title = TITLE
    doc.core_properties.author = BYLINE

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Écrit : {OUTPUT}")


if __name__ == "__main__":
    build()
